"""Safe, inspectable baseline parsers for the first real-data spike.

These providers are not claimed to be the final production parser combination.  They
extract native DOCX/XLSX structure and PDF text so actual samples can be inspected and
compared before introducing Docling/MinerU/RAGFlow routes.
"""

from __future__ import annotations

import datetime as dt
import io
import re
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from pypdf import PdfReader
import xlrd

from walkie_dokie.agents.security import validate_office_artifact

from .domain import ParseResult, ParsedBlock, normalize_text

_CLAUSE_RE = re.compile(r"^\s*(第[〇零一二三四五六七八九十百千万两0-9]+条(?:之[一二三四五六七八九十0-9]+)?)")
_NUMBERED_RE = re.compile(
    r"^\s*((?:\d+\.)+\d*|\d+[、.]|[一二三四五六七八九十百]+、|[（(][一二三四五六七八九十0-9]+[）)])"
)
_HEADING_STYLE_RE = re.compile(r"^(?:Heading|标题)\s*(\d+)$", re.IGNORECASE)
_MAX_PDF_BYTES = 100 * 1024 * 1024
_MAX_PDF_PAGES = 2_000
_MAX_WORKSHEET_CELLS = 2_000_000
_MAX_XLS_BYTES = 100 * 1024 * 1024
_COMPOUND_FILE_HEADER = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

_XLS_CELL_TYPE_NAMES = {
    xlrd.XL_CELL_EMPTY: "empty",
    xlrd.XL_CELL_TEXT: "text",
    xlrd.XL_CELL_NUMBER: "number",
    xlrd.XL_CELL_DATE: "date",
    xlrd.XL_CELL_BOOLEAN: "boolean",
    xlrd.XL_CELL_ERROR: "error",
    xlrd.XL_CELL_BLANK: "blank",
}


def _json_value(value: Any) -> Any:
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    if isinstance(value, (dt.date, dt.datetime, dt.time)):
        return value.isoformat()
    return str(value)


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _docx_audit(path: Path) -> dict[str, Any]:
    counts = {
        "insertions": 0,
        "deletions": 0,
        "hidden_text_runs": 0,
        "field_instructions": 0,
    }
    comments_present = False
    numbering_present = False
    with zipfile.ZipFile(path) as archive:
        names = {name.casefold() for name in archive.namelist()}
        comments_present = "word/comments.xml" in names
        numbering_present = "word/numbering.xml" in names
        for info in archive.infolist():
            lowered = info.filename.casefold()
            if not lowered.startswith("word/") or not lowered.endswith(".xml"):
                continue
            if info.file_size > 5 * 1024 * 1024:
                continue
            try:
                root = ElementTree.fromstring(archive.read(info))
            except ElementTree.ParseError:
                continue
            for element in root.iter():
                name = _local_name(element.tag)
                if name == "ins":
                    counts["insertions"] += 1
                elif name == "del":
                    counts["deletions"] += 1
                elif name in {"vanish", "webHidden"}:
                    counts["hidden_text_runs"] += 1
                elif name in {"instrText", "fldSimple"}:
                    counts["field_instructions"] += 1
    return {
        **counts,
        "comments_present": comments_present,
        "numbering_present": numbering_present,
    }


def _audit_warnings(audit: dict[str, Any]) -> list[str]:
    warnings: list[str] = []
    if audit["insertions"] or audit["deletions"]:
        warnings.append("DOCX 包含未决修订，发布前必须人工确认清稿政策")
    if audit["comments_present"]:
        warnings.append("DOCX 包含批注；批注未进入事实证据")
    if audit["hidden_text_runs"]:
        warnings.append("DOCX 包含隐藏文本；隐藏文本默认不进入检索")
    if audit["numbering_present"]:
        warnings.append("DOCX 使用自动编号；当前 baseline 仅保留审计标记，需用样例验证显示编号")
    return warnings


class NativeDocxParser:
    name = "native_docx"
    version = "0.1"

    def supports(self, path: Path) -> bool:
        return path.suffix.casefold() == ".docx"

    def parse(self, path: Path) -> ParseResult:
        validate_office_artifact(path, role="合同 ingestion 输入")
        audit = _docx_audit(path)
        document = Document(path)
        blocks: list[ParsedBlock] = []
        title_stack: list[str] = []
        ordinal = 0

        for paragraph_index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if not normalize_text(text):
                continue
            style_name = paragraph.style.name if paragraph.style is not None else ""
            heading_match = _HEADING_STYLE_RE.match(style_name or "")
            if heading_match:
                level = max(1, int(heading_match.group(1)))
                title_stack = title_stack[: level - 1]
                title_stack.append(normalize_text(text))

            clause_match = _CLAUSE_RE.match(text)
            numbered_match = _NUMBERED_RE.match(text)
            clause_ref = (
                clause_match.group(1)
                if clause_match
                else numbered_match.group(1) if numbered_match else None
            )
            structural_path = (*title_stack, clause_ref) if clause_ref else tuple(title_stack)
            blocks.append(
                ParsedBlock(
                    kind="heading" if heading_match else "clause" if clause_ref else "paragraph",
                    ordinal=ordinal,
                    text=text,
                    structural_path=tuple(part for part in structural_path if part),
                    title_path=tuple(title_stack),
                    clause_ref=clause_ref,
                    source_anchor={"type": "docx_paragraph", "paragraph_index": paragraph_index},
                    metadata={"style_name": style_name or ""},
                )
            )
            ordinal += 1

        for table_index, table in enumerate(document.tables):
            for row_index, row in enumerate(table.rows):
                cells = [normalize_text(cell.text) for cell in row.cells]
                if not any(cells):
                    continue
                cell_records = [
                    {"coordinate": f"R{row_index + 1}C{column_index + 1}", "text": value}
                    for column_index, value in enumerate(cells)
                ]
                blocks.append(
                    ParsedBlock(
                        kind="table_row",
                        ordinal=ordinal,
                        text=" | ".join(
                            f"C{column_index + 1}={value}"
                            for column_index, value in enumerate(cells)
                            if value
                        ),
                        structural_path=(*title_stack, f"表{table_index + 1}", f"行{row_index + 1}"),
                        title_path=tuple(title_stack),
                        source_anchor={
                            "type": "docx_table_row",
                            "table_index": table_index,
                            "row_index": row_index,
                            "cells": [cell["coordinate"] for cell in cell_records],
                        },
                        metadata={"cells": cell_records},
                    )
                )
                ordinal += 1

        return ParseResult(
            provider_name=self.name,
            provider_version=self.version,
            blocks=tuple(blocks),
            metadata={
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "audit": audit,
                "limitations": [
                    "baseline 尚未恢复 OOXML 自动编号的最终显示文本",
                    "baseline 的段落与表格未按 OOXML body 顺序交织",
                    "页码和 bbox 需要与正式 PDF 或受控渲染结果对齐",
                ],
            },
            warnings=tuple(_audit_warnings(audit)),
        )


def _merged_range_for(sheet, coordinate: str) -> str | None:
    for merged_range in sheet.merged_cells.ranges:
        if coordinate in merged_range:
            return str(merged_range)
    return None


class NativeXlsxParser:
    name = "native_xlsx"
    version = "0.1"

    def supports(self, path: Path) -> bool:
        return path.suffix.casefold() == ".xlsx"

    def parse(self, path: Path) -> ParseResult:
        validate_office_artifact(path, role="价目表 ingestion 输入")
        workbook = load_workbook(path, data_only=False, read_only=False, keep_links=False)
        cached_workbook = load_workbook(
            path, data_only=True, read_only=False, keep_links=False
        )
        blocks: list[ParsedBlock] = []
        warnings: list[str] = []
        ordinal = 0
        formula_count = 0
        hidden_sheet_count = 0

        for sheet in workbook.worksheets:
            if sheet.max_row * sheet.max_column > _MAX_WORKSHEET_CELLS:
                raise RuntimeError(
                    f"工作表 {sheet.title!r} 声明的单元格范围过大，需隔离人工检查"
                )
            cached_sheet = cached_workbook[sheet.title]
            if sheet.sheet_state != "visible":
                hidden_sheet_count += 1
            for row_index in range(1, sheet.max_row + 1):
                row_cells: list[dict[str, Any]] = []
                for column_index in range(1, sheet.max_column + 1):
                    cell = sheet.cell(row=row_index, column=column_index)
                    cached_cell = cached_sheet[cell.coordinate]
                    if cell.value is None and cached_cell.value is None:
                        continue
                    raw_value = _json_value(cell.value)
                    cached_value = _json_value(cached_cell.value)
                    is_formula = isinstance(cell.value, str) and cell.value.startswith("=")
                    formula_count += int(is_formula)
                    row_cells.append(
                        {
                            "coordinate": cell.coordinate,
                            "raw_value": raw_value,
                            "formula": raw_value if is_formula else None,
                            "cached_value": cached_value if is_formula else None,
                            "number_format": cell.number_format,
                            "merged_range": _merged_range_for(sheet, cell.coordinate),
                            "column_hidden": bool(sheet.column_dimensions[cell.column_letter].hidden),
                        }
                    )
                if not row_cells:
                    continue
                row_hidden = bool(sheet.row_dimensions[row_index].hidden)
                excluded_by_default = (
                    sheet.sheet_state != "visible"
                    or row_hidden
                )
                display_parts: list[str] = []
                for cell in row_cells:
                    display_value = (
                        cell["cached_value"]
                        if cell["formula"] is not None and cell["cached_value"] is not None
                        else cell["raw_value"]
                    )
                    display_parts.append(f"{cell['coordinate']}={display_value}")
                blocks.append(
                    ParsedBlock(
                        kind="spreadsheet_row",
                        ordinal=ordinal,
                        text=" | ".join(display_parts),
                        structural_path=(sheet.title, f"row:{row_index}"),
                        title_path=(sheet.title,),
                        source_anchor={
                            "type": "xlsx_row",
                            "sheet": sheet.title,
                            "row": row_index,
                            "cells": [cell["coordinate"] for cell in row_cells],
                        },
                        metadata={
                            "sheet_state": sheet.sheet_state,
                            "row_hidden": row_hidden,
                            "excluded_by_default": excluded_by_default,
                            "cells": row_cells,
                        },
                    )
                )
                ordinal += 1

        if formula_count:
            warnings.append(
                f"XLSX 包含 {formula_count} 个公式；当前仅保存公式与缓存值，尚未受控重算"
            )
        if hidden_sheet_count:
            warnings.append(
                f"XLSX 包含 {hidden_sheet_count} 个隐藏工作表；其内容默认不进入可信价格记录"
            )
        return ParseResult(
            provider_name=self.name,
            provider_version=self.version,
            blocks=tuple(blocks),
            metadata={
                "sheet_count": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames,
                "formula_count": formula_count,
                "hidden_sheet_count": hidden_sheet_count,
                "defined_names": [name for name in workbook.defined_names],
                "limitations": [
                    "尚未推断价目表表头或 MappingSpec",
                    "公式缓存值不等于经过受控重算的可信值",
                    "未解析的隐藏区域默认不能进入 Trusted 价格数据",
                ],
            },
            warnings=tuple(warnings),
        )


def _xls_cell_value(workbook, cell) -> Any:
    if cell.ctype == xlrd.XL_CELL_DATE:
        return xlrd.xldate_as_datetime(cell.value, workbook.datemode).isoformat()
    if cell.ctype == xlrd.XL_CELL_BOOLEAN:
        return bool(cell.value)
    return _json_value(cell.value)


def _xls_merged_range_for(sheet, row_index: int, column_index: int) -> str | None:
    for row_low, row_high, column_low, column_high in sheet.merged_cells:
        if row_low <= row_index < row_high and column_low <= column_index < column_high:
            start = f"{get_column_letter(column_low + 1)}{row_low + 1}"
            end = f"{get_column_letter(column_high)}{row_high}"
            return f"{start}:{end}"
    return None


class NativeXlsParser:
    """Read-only BIFF8 baseline that never claims to understand formula text."""

    name = "native_xls"
    version = "0.1"

    def supports(self, path: Path) -> bool:
        return path.suffix.casefold() == ".xls"

    def parse(self, path: Path) -> ParseResult:
        resolved = path.resolve()
        if not resolved.is_file():
            raise RuntimeError(f"XLS 不存在：{resolved}")
        if resolved.stat().st_size > _MAX_XLS_BYTES:
            raise RuntimeError("XLS 超过 100 MiB baseline 安全限制")
        with resolved.open("rb") as stream:
            if stream.read(8) != _COMPOUND_FILE_HEADER:
                raise RuntimeError("XLS 缺少 OLE Compound File 文件头")

        parser_log = io.StringIO()
        workbook = xlrd.open_workbook(
            resolved,
            on_demand=True,
            formatting_info=True,
            logfile=parser_log,
        )
        blocks: list[ParsedBlock] = []
        warnings = [
            "XLS 只能读取 Excel 保存的单元格值，无法完整保留公式文本；只可进入 Staging"
        ]
        ordinal = 0
        hidden_sheet_count = 0
        sheet_names = workbook.sheet_names()
        try:
            for sheet_name in sheet_names:
                sheet = workbook.sheet_by_name(sheet_name)
                if sheet.nrows * sheet.ncols > _MAX_WORKSHEET_CELLS:
                    raise RuntimeError(
                        f"工作表 {sheet_name!r} 声明的单元格范围过大，需隔离人工检查"
                    )
                sheet_hidden = bool(sheet.visibility)
                hidden_sheet_count += int(sheet_hidden)
                for row_index in range(sheet.nrows):
                    row_cells: list[dict[str, Any]] = []
                    row_hidden = bool(
                        sheet.rowinfo_map.get(row_index)
                        and sheet.rowinfo_map[row_index].hidden
                    )
                    for column_index in range(sheet.row_len(row_index)):
                        cell = sheet.cell(row_index, column_index)
                        if cell.ctype in {xlrd.XL_CELL_EMPTY, xlrd.XL_CELL_BLANK}:
                            continue
                        coordinate = (
                            f"{get_column_letter(column_index + 1)}{row_index + 1}"
                        )
                        column_info = sheet.colinfo_map.get(column_index)
                        row_cells.append(
                            {
                                "coordinate": coordinate,
                                "raw_value": _xls_cell_value(workbook, cell),
                                "formula": None,
                                "cached_value": None,
                                "cell_type": _XLS_CELL_TYPE_NAMES.get(
                                    cell.ctype, f"unknown:{cell.ctype}"
                                ),
                                "merged_range": _xls_merged_range_for(
                                    sheet, row_index, column_index
                                ),
                                "column_hidden": bool(
                                    column_info and column_info.hidden
                                ),
                                "formula_text_available": False,
                            }
                        )
                    if not row_cells:
                        continue
                    blocks.append(
                        ParsedBlock(
                            kind="spreadsheet_row",
                            ordinal=ordinal,
                            text=" | ".join(
                                f"{cell['coordinate']}={cell['raw_value']}"
                                for cell in row_cells
                            ),
                            structural_path=(sheet_name, f"row:{row_index + 1}"),
                            title_path=(sheet_name,),
                            source_anchor={
                                "type": "xls_row",
                                "sheet": sheet_name,
                                "row": row_index + 1,
                                "cells": [
                                    cell["coordinate"] for cell in row_cells
                                ],
                            },
                            metadata={
                                "sheet_hidden": sheet_hidden,
                                "row_hidden": row_hidden,
                                "excluded_by_default": sheet_hidden or row_hidden,
                                "cells": row_cells,
                                "formula_text_available": False,
                            },
                        )
                    )
                    ordinal += 1
                workbook.unload_sheet(sheet_name)
        finally:
            workbook.release_resources()

        parser_messages = tuple(
            line.strip() for line in parser_log.getvalue().splitlines() if line.strip()
        )
        if parser_messages:
            warnings.append(
                "XLS 公式解析器报告未识别函数；缓存值必须通过独立算术闭合验证"
            )
        if hidden_sheet_count:
            warnings.append(
                f"XLS 包含 {hidden_sheet_count} 个隐藏工作表；其内容默认不进入可信价格记录"
            )
        return ParseResult(
            provider_name=self.name,
            provider_version=self.version,
            blocks=tuple(blocks),
            metadata={
                "biff_version": workbook.biff_version,
                "encoding": workbook.encoding,
                "sheet_count": len(sheet_names),
                "sheet_names": sheet_names,
                "hidden_sheet_count": hidden_sheet_count,
                "formula_text_available": False,
                "parser_messages": list(parser_messages),
                "limitations": [
                    "XLS baseline 不保留公式文本，只读取 Excel 保存的单元格值",
                    "公式缓存值不等于经过受控重算的可信值",
                ],
            },
            warnings=tuple(warnings),
        )


def _pdf_segments(text: str) -> list[str]:
    lines = [normalize_text(line) for line in text.splitlines() if normalize_text(line)]
    if not lines:
        return []
    segments: list[str] = []
    current: list[str] = []
    current_chars = 0
    for line in lines:
        starts_structure = bool(_CLAUSE_RE.match(line) or _HEADING_STYLE_RE.match(line))
        if current and (starts_structure or current_chars + len(line) > 2_500):
            segments.append("\n".join(current))
            current = []
            current_chars = 0
        current.append(line)
        current_chars += len(line)
    if current:
        segments.append("\n".join(current))
    return segments


class TextPdfParser:
    name = "pypdf_text"
    version = "0.1"

    def supports(self, path: Path) -> bool:
        return path.suffix.casefold() == ".pdf"

    def parse(self, path: Path) -> ParseResult:
        resolved = path.resolve()
        if not resolved.is_file():
            raise RuntimeError(f"PDF 不存在：{resolved}")
        if resolved.stat().st_size > _MAX_PDF_BYTES:
            raise RuntimeError("PDF 超过 100 MiB baseline 安全限制")
        with resolved.open("rb") as stream:
            if stream.read(5) != b"%PDF-":
                raise RuntimeError("文件扩展名是 PDF，但缺少 PDF 文件头")
        reader = PdfReader(resolved, strict=True)
        if reader.is_encrypted:
            raise RuntimeError("PDF 已加密，当前 baseline 拒绝解析")
        if len(reader.pages) > _MAX_PDF_PAGES:
            raise RuntimeError("PDF 页数超过 baseline 安全限制")

        page_labels = list(getattr(reader, "page_labels", ()) or ())
        blocks: list[ParsedBlock] = []
        warnings: list[str] = []
        empty_pages: list[int] = []
        ordinal = 0
        for page_index, page in enumerate(reader.pages):
            page_number = page_index + 1
            try:
                page_text = page.extract_text() or ""
            except Exception as exc:
                warnings.append(f"PDF 第 {page_number} 页文字层提取失败：{type(exc).__name__}")
                empty_pages.append(page_number)
                continue
            segments = _pdf_segments(page_text)
            if not segments:
                empty_pages.append(page_number)
                continue
            for segment_index, text in enumerate(segments):
                clause_match = _CLAUSE_RE.match(text)
                clause_ref = clause_match.group(1) if clause_match else None
                blocks.append(
                    ParsedBlock(
                        kind="clause" if clause_ref else "pdf_text",
                        ordinal=ordinal,
                        text=text,
                        structural_path=(f"page:{page_number}", f"segment:{segment_index}"),
                        clause_ref=clause_ref,
                        page_physical=page_number,
                        page_printed=(
                            str(page_labels[page_index])
                            if page_index < len(page_labels)
                            else None
                        ),
                        source_anchor={
                            "type": "pdf_page_text",
                            "page": page_number,
                            "segment": segment_index,
                        },
                        metadata={"bbox": None, "text_layer": True},
                    )
                )
                ordinal += 1

        if empty_pages:
            warnings.append(
                f"PDF 有 {len(empty_pages)} 页没有可靠文字层，需要 OCR/layout parser：{empty_pages[:20]}"
            )
        return ParseResult(
            provider_name=self.name,
            provider_version=self.version,
            blocks=tuple(blocks),
            metadata={
                "page_count": len(reader.pages),
                "empty_text_pages": empty_pages,
                "text_coverage": (
                    (len(reader.pages) - len(empty_pages)) / len(reader.pages)
                    if reader.pages
                    else 0.0
                ),
                "limitations": [
                    "pypdf baseline 不提供可靠 bbox、阅读顺序或表格结构",
                    "无文字层页面必须路由到 MinerU/DeepDoc/OCR Provider",
                ],
            },
            warnings=tuple(warnings),
        )


def baseline_parser_registry():
    from .providers import ParserRegistry

    return ParserRegistry(
        (NativeDocxParser(), NativeXlsParser(), NativeXlsxParser(), TextPdfParser())
    )
