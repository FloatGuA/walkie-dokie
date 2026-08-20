"""生成 eval 用的最小合法 Office fixture。只需手动运行一次，产物提交进仓库。"""

from pathlib import Path

from docx import Document


def main() -> None:
    here = Path(__file__).parent
    doc = Document()
    doc.add_paragraph("这是 eval 输入样例文档。")
    doc.save(here / "simple.docx")
    out = Document()
    out.add_paragraph("这是 FakeExecutionAgent 的确定性产物。")
    out.save(here / "fake_output.docx")
    print("fixtures written")


if __name__ == "__main__":
    main()
