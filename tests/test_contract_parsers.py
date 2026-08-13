from io import BytesIO

from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from walkie_dokie.contract_intelligence.domain import stable_evidence_id
from walkie_dokie.contract_intelligence.parsers import (
    NativeDocxParser,
    NativeXlsxParser,
    TextPdfParser,
    baseline_parser_registry,
)


def _save_docx(path):
    document = Document()
    document.add_heading("第一章 服务范围", level=1)
    document.add_paragraph("第一条 服务价格为人民币 100 元。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "价格"
    table.cell(1, 0).text = "安装"
    table.cell(1, 1).text = "100元"
    document.save(path)


def test_native_docx_parser_preserves_clause_and_table_provenance(tmp_path):
    path = tmp_path / "contract.docx"
    _save_docx(path)

    result = NativeDocxParser().parse(path)

    clause = next(block for block in result.blocks if block.clause_ref == "第一条")
    table_row = next(block for block in result.blocks if block.kind == "table_row")
    assert clause.title_path == ("第一章 服务范围",)
    assert clause.source_anchor == {"type": "docx_paragraph", "paragraph_index": 1}
    assert table_row.source_anchor["cells"] == ["R1C1", "R1C2"]
    assert "价格" in table_row.text

    first_id = stable_evidence_id(
        document_version_id="v1", source_sha256="a" * 64, block=clause
    )
    second_id = stable_evidence_id(
        document_version_id="v1", source_sha256="a" * 64, block=clause
    )
    assert first_id == second_id
    assert first_id.startswith("ev_")


def test_native_xlsx_parser_keeps_coordinates_formula_and_hidden_state(tmp_path):
    path = tmp_path / "prices.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "价目表"
    sheet.append(["商品", "含税价", "数量"])
    sheet.append(["甲", 10, 2])
    sheet["D2"] = "=B2*C2"
    sheet.row_dimensions[3].hidden = True
    sheet["A3"] = "隐藏价格"
    hidden = workbook.create_sheet("内部参数")
    hidden.sheet_state = "hidden"
    hidden["A1"] = "不可默认发布"
    workbook.save(path)

    result = NativeXlsxParser().parse(path)

    price_row = next(
        block
        for block in result.blocks
        if block.source_anchor.get("sheet") == "价目表"
        and block.source_anchor.get("row") == 2
    )
    formula_cell = next(
        cell for cell in price_row.metadata["cells"] if cell["coordinate"] == "D2"
    )
    hidden_row = next(
        block
        for block in result.blocks
        if block.source_anchor.get("sheet") == "价目表"
        and block.source_anchor.get("row") == 3
    )
    assert formula_cell["formula"] == "=B2*C2"
    assert hidden_row.metadata["excluded_by_default"] is True
    assert result.metadata["hidden_sheet_count"] == 1
    assert any("受控重算" in warning for warning in result.warnings)


def test_pdf_without_text_layer_is_explicitly_flagged(tmp_path):
    path = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    with path.open("wb") as output:
        writer.write(output)

    result = TextPdfParser().parse(path)

    assert result.blocks == ()
    assert result.metadata["empty_text_pages"] == [1]
    assert any("OCR" in warning for warning in result.warnings)


def test_baseline_registry_routes_by_explicit_file_type(tmp_path):
    path = tmp_path / "contract.docx"
    _save_docx(path)
    provider = baseline_parser_registry().resolve(path)
    assert provider.name == "native_docx"
