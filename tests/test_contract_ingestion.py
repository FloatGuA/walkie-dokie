from io import BytesIO

import pytest
from django.contrib import admin
from django.contrib.auth import get_user_model
from django.core.exceptions import ValidationError
from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from django.test import RequestFactory
from docx import Document as WordDocument

from walkie_dokie.contract_intelligence.ingestion import run_parser_for_source
from walkie_dokie.contract_intelligence.models import (
    Document,
    DocumentVersion,
    EvidenceUnit,
    KnowledgeProject,
    ParserRun,
    RetrievalTrace,
    RetrievalUnit,
    SourceFile,
    ProjectMembership,
)


def _docx_bytes() -> bytes:
    buffer = BytesIO()
    document = WordDocument()
    document.add_heading("费用", level=1)
    document.add_paragraph("第一条 服务费为人民币 100 元。")
    document.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def contract_version(db):
    project = KnowledgeProject.objects.create(name="测试项目", slug="test-project")
    document = Document.objects.create(
        project=project, name="服务合同", kind=Document.Kind.CONTRACT
    )
    return DocumentVersion.objects.create(document=document, version_label="v1")


@pytest.mark.django_db
def test_ingestion_persists_run_evidence_and_retrieval_units(
    contract_version, settings, tmp_path
):
    settings.MEDIA_ROOT = tmp_path / "media"
    source = SourceFile(
        document_version=contract_version,
        role=SourceFile.Role.STRUCTURED_SOURCE,
        file=SimpleUploadedFile(
            "contract.docx",
            _docx_bytes(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
    )
    source.full_clean()
    source.save()

    run = run_parser_for_source(source.pk, raise_errors=True)

    contract_version.refresh_from_db()
    assert run.status == ParserRun.Status.SUCCEEDED
    assert contract_version.status == DocumentVersion.Status.REVIEW
    assert EvidenceUnit.objects.filter(parser_run=run).count() == 2
    assert RetrievalUnit.objects.filter(evidence__parser_run=run).count() == 2
    clause = EvidenceUnit.objects.get(parser_run=run, clause_ref="第一条")
    assert clause.source_anchor["paragraph_index"] == 1
    assert clause.evidence_id.startswith("ev_")
    assert source.sha256 and source.size_bytes == len(_docx_bytes())


@pytest.mark.django_db
def test_source_file_and_published_version_are_immutable(
    contract_version, settings, tmp_path
):
    settings.MEDIA_ROOT = tmp_path / "media"
    source = SourceFile.objects.create(
        document_version=contract_version,
        role=SourceFile.Role.STRUCTURED_SOURCE,
        file=SimpleUploadedFile("contract.docx", _docx_bytes()),
    )
    source.mime_type = "changed/type"
    with pytest.raises(ValidationError, match="不允许原地修改"):
        source.save()

    DocumentVersion.objects.filter(pk=contract_version.pk).update(
        status=DocumentVersion.Status.PUBLISHED
    )
    contract_version.refresh_from_db()
    contract_version.notes = "attempted overwrite"
    with pytest.raises(ValidationError, match="不允许原地修改"):
        contract_version.save()


@pytest.mark.django_db
def test_admin_inspection_page_exposes_exact_evidence(
    client, contract_version, settings, tmp_path
):
    settings.MEDIA_ROOT = tmp_path / "media"
    user = get_user_model().objects.create_superuser(
        username="admin", email="admin@example.com", password="safe-test-password"
    )
    source = SourceFile.objects.create(
        document_version=contract_version,
        role=SourceFile.Role.STRUCTURED_SOURCE,
        file=SimpleUploadedFile("contract.docx", _docx_bytes()),
    )
    run_parser_for_source(source.pk, raise_errors=True)
    client.force_login(user)

    response = client.get(
        reverse("admin:contract_document_version_inspect", args=(contract_version.pk,))
    )

    assert response.status_code == 200
    page = response.content.decode("utf-8")
    assert "第一条 服务费为人民币 100 元" in page
    assert "Evidence Units" in page

    retrieval_response = client.post(
        reverse("admin:contract_document_version_inspect", args=(contract_version.pk,)),
        {"retrieval_q": "服务费多少钱", "top_k": "2"},
    )
    assert retrieval_response.status_code == 200
    retrieval_page = retrieval_response.content.decode("utf-8")
    assert "local_chinese_bm25" in retrieval_page
    assert "stage scores" in retrieval_page
    trace = RetrievalTrace.objects.get()
    assert trace.query == "服务费多少钱"
    assert trace.candidates[0]["evidence_id"].startswith("ev_")


@pytest.mark.django_db
def test_non_superuser_admin_queryset_is_project_scoped():
    user = get_user_model().objects.create_user(username="staff", is_staff=True)
    allowed = KnowledgeProject.objects.create(name="可访问", slug="allowed")
    KnowledgeProject.objects.create(name="不可访问", slug="denied")
    ProjectMembership.objects.create(project=allowed, user=user)
    request = RequestFactory().get("/admin/")
    request.user = user

    model_admin = admin.site._registry[KnowledgeProject]
    visible = list(model_admin.get_queryset(request))

    assert visible == [allowed]
